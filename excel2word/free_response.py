from openpyxl import load_workbook
workbook = load_workbook(filename="Survey.xlsx") 

free_responses={}
for sheet in workbook.worksheets:
    for i in range(1,50):
        s=sheet[f"A{i}"].value
        if s =='Respondent ID': # if the question is free-response
            res=[]
            #for cell in sheet[f'C{i+1}:C1048576']:
            for value in sheet.iter_rows(min_row=i+1,min_col=3, max_col=3,values_only=True):
                if (value[0]!=None):
                    res.append(value[0]) # add res to that single question 
            #print(res)
            free_responses[sheet.title]=res  #add list of res to one question to the big list
            break

import docx

doc = docx.Document()
for key,values in free_responses.items():
    doc.add_paragraph(key)
    for res in values:
        doc.add_paragraph(res,style='List Bullet 2')
doc.save('output.docx')