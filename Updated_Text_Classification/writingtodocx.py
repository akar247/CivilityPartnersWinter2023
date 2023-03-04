import docx
import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification, \
    pipeline

import pandas as pd
import os
import numpy as np
import re


doc = docx.Document('Copy of Interview Notes.docx')

lst = []
for para in doc.paragraphs:
    lst.append(para.text)

df = pd.DataFrame(data=lst, columns=['text'])
df = df[df['text'] != '']  # dropping empty text

inputs = lst[0]
regex = '[A-Z]{1}[a-z]* - .*'
headers = df[df['text'].str.contains(regex)]

index = [lst[0]]
text = []
for i in range(1, len(lst) - 1):
    string = lst[i]
    if lst[i - 1] == '' and lst[i + 1] == '' and lst[i] != '':
        index.append(string)
    else:
        text.append(string)

headers = list(set(headers['text'].tolist()).intersection(index))
text_df = df[~df['text'].isin(headers)]

tokenizer = AutoTokenizer.from_pretrained("DhruvK0/test_trainer")
model = AutoModelForSequenceClassification.from_pretrained(
    "DhruvK0/test_trainer")
classifier = pipeline("sentiment-analysis", model=model, tokenizer=tokenizer)


def give_sentiment(text):
    return classifier(text)[0]['label']

print('Analyzing sentiment')

text_df['Output'] = text_df['text'].apply(give_sentiment)

print('Sentiment analyzed')

text_df['Output Sign'] = text_df['Output'].transform(lambda sentiment:
                                                     '+' if sentiment == 'POSITIVE' else '-')

output_doc = docx.Document()


def add_to_docx(row):
    add_text = row['Output Sign'] + row['text']
    output_p = output_doc.add_paragraph(add_text)
    output_p.style = 'List Bullet'


text_df.apply(add_to_docx, axis=1)

output_doc.save('test_output.docx')

print('Saved output document to drive')
