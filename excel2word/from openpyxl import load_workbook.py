from openpyxl import load_workbook
workbook = load_workbook(filename="Survey.xlsx") 

multiple_choices={}
for sheet in workbook.worksheets:
    for i in range(1,50):
        s=sheet[f"A{i}"].value
        if s =='Answer Choices': # if the question is multiple choice
            res=[]
            #for cell in sheet[f'C{i+1}:C1048576']:
            for value in sheet.iter_rows(min_row=i+1,min_col=3, max_col=3,values_only=True):
                if (value[0]!=None):
                    res.append(value[0]) # add res to that single question 
            #print(res)
            multiple_choices[sheet.title]=res  #add list of res to one question to the big list
            break

import docx

from openpyxl import load_workbook
import os

workbook = load_workbook(filename="Survey.xlsx") 
Sheet_Names = [workbook.sheetnames]
#workbook.sheetnames
#workbook.worksheets
import docx

free_responses={}
for sheet in workbook.worksheets:
    lis=[]
    freechart=[]
    free_responses[sheet.title] = lis
    for i in range(1,50):
        s=sheet[f"A{i}"].value
        res=[]
        if s =='Respondent ID': # if the question is free-response
            

            #for cell in sheet[f'C{i+1}:C1048576']:
            for value in sheet.iter_rows(min_row=i+1,min_col=3, max_col=3,values_only=True):
                if (value[0]!=None):
                    res.append(value[0]) # add res to that single question 
            #print(res)
            free_responses[sheet.title]=res  #add list of res to one question to the big list
            break
            
        
        
            

#free_responses['Question 25']
from openpyxl.chart import BarChart, Series, Reference
from openpyxl.chart.label import DataLabelList
wb = load_workbook(filename="Survey.xlsx") 
ws = wb.active

for ws in wb.worksheets:
    # Data for plotting

    

    

    first_empty_cell = ws.max_row+1

    values = Reference(ws, min_col= 2, min_row=3, max_col = 2, max_row = first_empty_cell - 1)
    cats = Reference(ws, min_col=1, max_col=1, min_row=4, max_row = first_empty_cell - 1)

    # Create object of BarChart class
    chart = BarChart()
    chart.height = 10
    chart.width = 15
    chart.add_data(values, titles_from_data=True)
    chart.set_categories(cats)
    # set the title of the chart
    chart.title = ws["A2"].value

    # the top-left corner of the chart
    # is anchored to cell F2 .
    chart.varyColors = "000F0FFF"
    ws.add_chart(chart,"H2")
    

# save the file 
wb.save("test.xlsx")   ###check the test result at new created "test.xlsx" file 

doc = docx.Document()

for key,values in free_responses.items():
    length_key = len(free_responses[key])
    ##print(length_key)
    doc.add_paragraph(key)
    ##if question is free response only OR multiple choice free response questions 
    if length_key > 1: 
       
        ##if question has both free response questions AND a chart then add the chart
        if os.path.exists(f'Images/Survey{key}.png'):
            doc.add_picture(f'Images/Survey{key}.png')
        ##adds free response answers in bullet format 
        for res in values:
            doc.add_paragraph(res,style='List Bullet 2')
    ##if question does not contain any free response responses 
    if length_key < 1:
        doc.add_picture(f'Images/Survey{key}.png')
doc.save('output.docx')