from tkinter import *
from tkinter import filedialog
from tkinter import font
from tkinter.ttk import Progressbar
from tkinter.ttk import Scrollbar
from openpyxl import load_workbook
import docx
workbook = load_workbook(filename="Survey.xlsx") 


#main window
window = Tk()
window.title("Excel2Word")
window.geometry('600x300')
window.config(bg='#ffffff')

bg_color = 'white'
title_color = 'purple'

# frames
frame1 = Frame(bg=title_color) # holds the title
frame1.pack(fill='x')
frame6 = Frame(bg=bg_color) # holds the interactive message
frame6.pack()
frame2 = Frame(bg=bg_color) # holds the choose file, preview, and start buttons
frame2.pack()
frame3 = Frame(bg=bg_color) # holds the progress bar
frame3.pack(fill='x')

#Global Variables
xfile = ''

# handler for 'Choose File' button click event
def open_xfile():
    global xfile
    xfile = filedialog.askopenfilename(title='open a file',filetypes=(('excel files','*.xls'),('excel files','*.xlsx')))
    lbl_filename.config(text=xfile[:25],fg='grey')
    extract_FR()

def extract_images():
    global xfile
    # app = Dispatch("Excel.Application")
    # workbook = app.Workbooks.Open(Filename=xfile)

    # i = 1
    for sheet in workbook.Worksheets:
        for chartObject in sheet.ChartObjects():
            print(sheet.Name + ':' + chartObject.Name)
            chartObject.Chart.Export(xfile + sheet.Name + ".png")
            # i += 1

    workbook.Close(SaveChanges=False, Filename=xfile)
    lbl_message.configure(text='Extracted Images')
    # extract_FR()

def extract_FR():
    global xfile
    workbook = load_workbook(filename=xfile) 
    
    free_responses={}
    for sheet in workbook.worksheets:
        for i in range(1,50):
            s=sheet[f"A{i}"].value
            if s =='Respondent ID': # if the question is free-response
                res=[]
                #for cell in sheet[f'C{i+1}:C1048576']:
                for value in sheet.iter_rows(min_row=i+1,min_col=3, max_col=3,values_only=True):
                    if (value[0]!=None):
                        res.append(value[0]) # add res to that single question 
                #print(res)
                free_responses[sheet.title]=res  #add list of res to one question to the big list
                break
                
    doc = docx.Document()
    for key,values in free_responses.items():
        doc.add_paragraph(key)
        for res in values:
            doc.add_paragraph(res,style='List Bullet 2')
    doc.save(xfile[:-5]+'Grabbed.docx')
    lbl_message.configure(text='Extracted Text')




#BELOW IS ALL UI
# create widgets (labels, buttons, progress bar, scroll bar)
lbl_title = Label(master=frame1,font=('Arial',25,'bold'),text='Excel to Word Conversion',fg='white',bg=title_color)
lbl_filename = Label(master=frame2,text='No File Chosen',fg='grey',width=20,anchor='w')
lbl_message = Label(master=frame6,text='*Please upload a file',bg='white', fg=title_color)

btn_choosefile = Button(master=frame2,text='Choose File',command=open_xfile, fg = "white", font = "Future 10", bg=title_color)

# pack widgets
lbl_title.pack(padx=20,pady=20)
lbl_filename.grid(sticky='w',row=0,column=1,padx=10,pady=10)
lbl_message.pack(pady=5)

btn_choosefile.grid(sticky='w',row=0,column=0,padx=10,pady=10)

# execute
window.mainloop() 
