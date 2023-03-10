from tkinter import *
import tkinter as tk
from tkinter import filedialog
from tkinter import font
from tkinter.ttk import Progressbar
from tkinter.ttk import Scrollbar
from os.path import exists
import csv
import ntpath
import pandas as pd
import docx
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem.snowball import SnowballStemmer
import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification, \
    pipeline
import re

import time

# UI

bg_button = 'sea green'
msg = 'blue'
bot_bg = 'medium aquamarine'
txt_bg = 'LavenderBlush4'
txt_color = 'gold'

# MAIN HOME

home = Tk()
home.title("Civility Partners Statement Classification")
home.geometry('600x300')
home.config(bg=bot_bg)

# FRAMES

frame1 = Frame(bg=bg_button)  # holds the title
frame1.pack(fill='x')
frame6 = Frame(bg=bot_bg)  # holds the interactive message
frame6.pack()
frame2 = Frame(bg=bot_bg)  # holds the choose file, preview, and start buttons
frame2.pack()
frame3 = Frame(bg=bot_bg)  # holds the progress bar
frame3.pack(fill='none')
frame4 = Frame(bg=bot_bg)
frame4.pack(fill='none')

# ENTRY TEXT

class_var = tk.StringVar()
c1_kw = tk.StringVar()
c2_kw = tk.StringVar()
c3_kw = tk.StringVar()
c4_kw = tk.StringVar()
c5_kw = tk.StringVar()

# END OF UI

# CONSTANT VARIABLES

stemmer = SnowballStemmer(language="english")
stopwords = set(stopwords.words('english'))
file = ''
error = False
tokenizer = AutoTokenizer.from_pretrained("DhruvK0/test_trainer")
model = AutoModelForSequenceClassification.from_pretrained(
    "DhruvK0/test_trainer")
classifier = pipeline("sentiment-analysis", model=model, tokenizer=tokenizer)

# Universal Vars
# list of classes they want
# classes = [] Not sure why it is defined here when it is redefined later
# list of lists of keys words for each class. Outside list should have same size as classes list
# kws = [] Not sure why it is defined here when it is redefined later
# Storing all responses
data = {}

# DEFAULT TESTING
job = ['job', 'task', 'responsibility', 'promotion', 'career', 'pay', 'salary',
       'compensation', 'possibilities', 'paid', 'flexibility', 'efficiency',
       'advancement', 'growth']
engagement = ['challenging', 'engagement', 'environment', "integrity",
              "meaningful", "timeline", "motivated", "goal", "focus"]
communication = ['communication', 'internal', 'organize', 'organization',
                 'goal', 'vision', 'communicate', 'collaborate', 'team',
                 'performance', 'manager', 'teach', 'ask', 'feedback',
                 'supervisor', 'transparent']
inclusion = ['inclusive', 'diversity', 'enforce', 'harass', 'comfort',
             'equality', 'behavior', 'hire', 'retain', 'talent', 'recruitment']
relationship = ['management', 'relation', 'relationship', 'leadership',
                'transparency', 'cohesiveness', 'accessible', 'straining',
                'collaborative', 'acceptance', 'supportive', 'respectful',
                'appreciative', 'friendly', 'approachable', 'honest',
                'relaxed', 'casual', 'toxic', 'dynamic', 'culture', 'moral',
                'trust']

classes = ['Job', 'Engagement', 'Communication', 'Inclusion', 'Relationship']
kws = [job, engagement, communication, inclusion, relationship]


# Thresholds output file code, takes file name as a string
def thresholds(filename):
    statements = {}

    def read_excel_questions(fp):
        sheets = pd.ExcelFile(fp).sheet_names
        q_data = []

        for sheet in sheets:
            q_data.append(
                (sheet, pd.read_excel(fp, sheet_name=sheet, header=2)))
            statements[sheet] = pd.read_excel(fp, sheet_name=sheet).iloc[0, 0]
        return q_data

    def mc_data(dfs):
        col_check = ['Answer Choices', 'Responses', 'Unnamed: 2']
        tholds = {'Negative': [], 'Neutral': [], 'Positive': []}
        for i in range(len(dfs)):
            sheet = dfs[i][0]
            df = dfs[i][1]
            if list(df.columns) == col_check:
                neg = df['Responses'].iloc[:2].sum()
                pos = df['Responses'].iloc[2:4].sum()

                if neg >= .25:
                    tholds['Negative'].append(
                        (sheet, statements[sheet], round(neg, 2)))
                elif pos >= .90:
                    tholds['Positive'].append(
                        (sheet, statements[sheet], round(pos, 2)))
                else:
                    tholds['Neutral'].append(sheet)

        return tholds

    def write_txt(dict):
        with open(file.split('.')[0] + '_thresholds.txt', 'w') as f:
            f.write('Negatives:\n')
            for qt in dict['Negative']:
                f.write('\t' + str(qt[0]) + '\n')
                f.write(
                    '\t\t' + str(qt[1]) + " ({:.0%}) ".format(qt[2]) + '\n\n')
            f.write('\n\n')
            f.write('Positives:\n')
            for qt in dict['Positive']:
                f.write('\t' + str(qt[0]) + '\n')
                f.write(
                    '\t\t' + str(qt[1]) + " ({:.0%}) ".format(qt[2]) + '\n\n')

    q_data = read_excel_questions(filename)
    # print(q_data[:5], '1')
    mc_tholds = mc_data(q_data)
    write_txt(mc_tholds)


def grab_data():
    global file

    if file.lower().endswith(('.xls', '.xlsx')):
        df = pd.read_excel(file, None)
        key = list(df)
        for i in range(len(df)):
            index = key[i]
            if df[index]["Unnamed: 2"].dtype == 'O':
                data[index] = df[index]["Unnamed: 2"].dropna()[1:]
                data[index] = [j for j in data[index] if
                               not isinstance(j, int)]
    else:
        doc = docx.Document(file)
        lst = []
        for para in doc.paragraphs:
            lst.append(para.text)

        doc_df = pd.DataFrame(data=lst, columns=['text'])
        doc_df = doc_df[doc_df['text'] != '']  # dropping empty text

        inputs = lst[0]
        regex = '[A-Z]{1}[a-z]* - .*'
        headers = doc_df[doc_df['text'].str.contains(regex)]

        index = [lst[0]]
        text = []
        for i in range(1, len(lst) - 1):
            string = lst[i]
            if lst[i - 1] == '' and lst[i + 1] == '' and lst[i] != '':
                index.append(string)
            else:
                text.append(string)
        headers = list(set(headers['text'].tolist()).intersection(index))
        doc_df = doc_df[~doc_df['text'].isin(headers)]
        return doc_df

    return data


def sentiment_score(response):
    """
    Sentiment analysis between positive or negative
    """

    return classifier(response)[0]['label']


def class_choice():
    global class_var, c1_kw, c2_kw, c3_kw, c4_kw, c5_kw, classes, kws, file, error
    reset_screen()
    if not exists(file):
        lbl_message.configure(text='Invalid Filepath')
        file = 'None'
        error = True
        return home_screen()
    try:
        if file.lower().endswith(('.xls', '.xlsx')):
            pd.read_excel(file)
        else:
            docx.Document(file)
    except:
        error = True
        lbl_message.configure(
            text='FileOpen Error: Make sure file is not open')
        return home_screen()
    try:
        if file.lower().endswith(('.xls', '.xlsx')):
            thresholds(file)
    except:
        error = True
        lbl_message.configure(
            text='Thresholds Error: Make sure file is formatted correctly')
        return home_screen()
    try:
        btn_cont.place_forget()
    except:
        pass
    lbl_title.configure(text='Keyword Edits')
    lbl_message.configure(
        text='If required, edit keywords with comma separation')
    lbl_c1.grid(sticky='w', row=50, column=0, padx=5, pady=5)
    lbl_c2.grid(sticky='w', row=51, column=0, padx=5, pady=5)
    lbl_c3.grid(sticky='w', row=52, column=0, padx=5, pady=5)
    lbl_c4.grid(sticky='w', row=53, column=0, padx=5, pady=5)
    lbl_c5.grid(sticky='w', row=54, column=0, padx=5, pady=5)
    entry_kw1.grid(sticky='w', row=50, column=2, padx=5, pady=5)
    entry_kw2.grid(sticky='w', row=51, column=2, padx=5, pady=5)
    entry_kw3.grid(sticky='w', row=52, column=2, padx=5, pady=5)
    entry_kw4.grid(sticky='w', row=53, column=2, padx=5, pady=5)
    entry_kw5.grid(sticky='w', row=54, column=2, padx=5, pady=5)
    entry_kw1.insert(0, ', '.join(kws[0]))
    entry_kw2.insert(0, ', '.join(kws[1]))
    entry_kw3.insert(0, ', '.join(kws[2]))
    entry_kw4.insert(0, ', '.join(kws[3]))
    entry_kw5.insert(0, ', '.join(kws[4]))
    btn_classify.place(x=475, y=175)


def classify():
    global classes, kws, stemmer, stopwords, data, c1_kw, c2_kw, c3_kw, c4_kw, data
    data = grab_data()
    lbl_message.configure(text='Classified Responses. Word Document Created!')
    lbl_title.configure(text='Classification')

    try:
        btn_classify.place_forget()
    except:
        pass

    if file.lower().endswith(('.xls', '.xlsx')):
        classified = {'Miscellaneous': []}
        kws[0] = c1_kw.get().split(', ')
        kws[1] = c2_kw.get().split(', ')
        kws[2] = c3_kw.get().split(', ')
        kws[3] = c4_kw.get().split(', ')
        kws[4] = c5_kw.get().split(', ')

        for i in range(len(kws)):
            kws[i] = [stemmer.stem(word) for word in kws[i]]

        for lst in range(len(kws)):
            for i in range(len(kws[lst])):
                # Apply Stemmer Filter on all single word keywords
                # Cannot stem phrases
                if ' ' not in kws[lst][i].strip():
                    kws[lst][i] = stemmer.stem(kws[lst][i])
            # create empty dictionary of classes key values
            classified[classes[lst]] = []

        for question in data:
            # formatting response
            responses = data[question]

            for response in responses:
                response_tok = word_tokenize(response)
                clean_text = [stemmer.stem(word) for word in response_tok if
                              word.isalpha() and word not in stopwords]

                sims = [0] * len(kws)

                for i in range(len(kws)):
                    sims[i] = len(set(kws[i]) & set(clean_text))

                if max(sims) == 0:
                    classified['Miscellaneous'].append(response)
                else:
                    temp_cat = classes[sims.index(max(sims))]
                    classified[temp_cat].append(response)
        doc = docx.Document()
        for c in classified:
            doc.add_paragraph(c)
            # classified[c] = sentiment_score(classified[c])
            df25 = pd.DataFrame(classified[c], columns=['response'])
            # df25.iloc[0]
            sentiment_score(df25['response'].iloc[1])
            df25['sentiment'] = df25['response'].apply(
                lambda x: sentiment_score(x[:512]))
            df25 = df25.sort_values('sentiment', ascending=False)
            for x, y in zip(df25['response'], df25['sentiment']):
                if y == "POSITIVE":
                    doc.add_paragraph('+' + x, style='List Bullet 2')
                elif y == "NEGATIVE":
                    doc.add_paragraph('-' + x, style='List Bullet 2')
        doc.save(file.split('.')[0] + '_output.docx')

    else:
        data['Output'] = data['text'].apply(sentiment_score)
        data['Output Sign'] = data['Output'].transform(lambda sentiment:
                                                             '+' if sentiment == 'POSITIVE' else '-')
        doc = docx.Document()

        def add_to_docx(row):
            add_text = row['Output Sign'] + row['text']
            output_p = doc.add_paragraph(add_text)
            output_p.style = 'List Bullet'

        data.apply(add_to_docx, axis=1)

        doc.save(file.split('.')[0] + '_output.docx')

        return

    # with open('practice.txt', 'w') as f:
    #     for key in classified:
    #         f.write(key.upper()+'\n')
    #         for val in classified[key]:
    #             f.write('\t\t'+val+'\n')
    #         f.write('\n\n')
    # print(classified)
    return classified


def open_file():
    """
        Handler for opening excel and docx files
    """
    global file

    file = filedialog.askopenfilename(title='open a file', filetypes=(
        ('excel files', '*.xls'), ('excel files', '*.xlsx'),
        ('docx files', '*.docx'), ('docx files', '*.doc')))

    if not exists(file):
        error_msg = 'Invalid Filepath'
        file = 'None'
        return home_screen()

    if len(file) > 30:
        lbl_filename.configure(text=file[:27] + '...')
    else:
        lbl_filename.configure(text=file)

    # UI Features


lbl_title = Label(master=frame1, font=('Arial', 25, 'bold'),
                  text='Transaction Categorization', fg='white', bg=bg_button)
lbl_filename = Label(master=frame2, text='No File Chosen', fg='grey', width=30,
                     anchor='w')
lbl_message = Label(master=frame6, text='*Please select an Option*', bg=bot_bg,
                    fg=msg)
lbl_c1 = Label(master=frame2, text=classes[0], fg=txt_color, width=15,
               anchor='w', bg=txt_bg, font=('Stencil Std', 10, 'bold'))
lbl_c2 = Label(master=frame2, text=classes[1], fg=txt_color, width=15,
               anchor='w', bg=txt_bg, font=('Stencil Std', 10, 'bold'))
lbl_c3 = Label(master=frame2, text=classes[2], fg=txt_color, width=15,
               anchor='w', bg=txt_bg, font=('Stencil Std', 10, 'bold'))
lbl_c4 = Label(master=frame2, text=classes[3], fg=txt_color, width=15,
               anchor='w', bg=txt_bg, font=('Stencil Std', 10, 'bold'))
lbl_c5 = Label(master=frame2, text=classes[4], fg=txt_color, width=15,
               anchor='w', bg=txt_bg, font=('Stencil Std', 10, 'bold'))

btn_choosefile = Button(master=frame2, text='Choose File', command=open_file,
                        fg="white", font="Future 10", bg=bg_button)
btn_cont = Button(master=home, text='Continue?', command=class_choice,
                  fg="white", font="Future 10", bg=bg_button, height=1,
                  width=10)
btn_classify = Button(master=home, text='Classify?',
                      command=lambda: [reset_screen(), classify()], fg="white",
                      font="Future 10", bg=bg_button, height=1, width=10)

entry_classes = tk.Entry(frame2, textvariable=class_var,
                         font=('calibre', 10, 'normal'), justify='center',
                         bg='white', fg='black')
entry_kw1 = tk.Entry(frame2, textvariable=c1_kw,
                     font=('calibre', 10, 'normal'), justify='center',
                     bg='white', fg='black')
entry_kw2 = tk.Entry(frame2, textvariable=c2_kw,
                     font=('calibre', 10, 'normal'), justify='center',
                     bg='white', fg='black')
entry_kw3 = tk.Entry(frame2, textvariable=c3_kw,
                     font=('calibre', 10, 'normal'), justify='center',
                     bg='white', fg='black')
entry_kw4 = tk.Entry(frame2, textvariable=c4_kw,
                     font=('calibre', 10, 'normal'), justify='center',
                     bg='white', fg='black')
entry_kw5 = tk.Entry(frame2, textvariable=c5_kw,
                     font=('calibre', 10, 'normal'), justify='center',
                     bg='white', fg='black')


def reset_screen():
    """
        Resets UI screen by removing any widgets present from previous screen
    """
    for widget in frame2.winfo_children():
        try:
            widget.grid_remove()
        except:
            widget.place_forget()
    for widget in frame3.winfo_children():
        try:
            widget.grid_remove()
        except:
            widget.place_forget()


def home_screen():
    """
        starting screen where original decision is made
        will return to this screen whenever finished or illegal action made
    """
    global error
    reset_screen()
    try:
        btn_back.place_forget()
    except:
        pass
    if not error:
        lbl_message.configure(text='*Please Choose Excel or Docx File*')
    # pack widgets
    lbl_title.configure(text='Statement Classification')
    lbl_title.pack(padx=20, pady=20)
    lbl_message.pack(pady=5)
    lbl_filename.grid(sticky='w', row=0, column=1, padx=10, pady=10)
    btn_choosefile.grid(sticky='w', row=0, column=0, padx=10, pady=10)
    btn_cont.place(x=250, y=220)


btn_back = Button(master=home, text='Back', command=home_screen, fg="white",
                  font="Future 10", bg=bg_button, height=2, width=10)

home_screen()
home.mainloop()
